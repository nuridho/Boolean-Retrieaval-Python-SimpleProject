import tkinter
import tkinter.messagebox
import customtkinter
import os
import docx
from PyPDF2 import PdfReader
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
import string
import sys

customtkinter.set_appearance_mode("System")
customtkinter.set_default_color_theme("blue")


class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()
        self.popup = None
        # configure window
        self.title("Program Pencarian menggunakan Boolean Retrieval")
        # self.geometry(f"{1200}x{650}")
        # Mendapatkan ukuran layar
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        window_width = 1200
        window_height = 650

        # Menghitung posisi untuk window berada di tengah
        x_position = (screen_width - window_width) // 2
        y_position = (screen_height - window_height) // 2

        self.geometry(f"{window_width}x{window_height}+{x_position}+{y_position}")

        # configure grid layout (3x3)
        # weight disini itu bobot.. jadi kolom 1 itu lebih dominan di banding  kolom 2 (1>0)
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(2, weight=0)
        self.grid_rowconfigure((0, 1, 2), weight=1)

        # create sidebar frame with widgets
        self.sidebar_frame = customtkinter.CTkFrame(self, width=140, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, rowspan=3, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(3, weight=1)
        self.logo_label = customtkinter.CTkLabel(self.sidebar_frame, text="BooleanRetrieval",
                                                 font=customtkinter.CTkFont(size=20, weight="bold"))
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 10))
        self.sidebarMateri = customtkinter.CTkButton(self.sidebar_frame)
        self.sidebarMateri.grid(row=1, column=0, padx=20, pady=10)
        self.appearance_mode_label = customtkinter.CTkLabel(self.sidebar_frame, text="Appearance Mode:", anchor="w")
        self.appearance_mode_label.grid(row=5, column=0, padx=20, pady=(10, 0))
        self.appearance_mode_optionemenu = customtkinter.CTkOptionMenu(self.sidebar_frame,
                                                                       values=["Light", "Dark", "System"],
                                                                       command=self.change_appearance_mode_event)
        self.appearance_mode_optionemenu.grid(row=6, column=0, padx=20, pady=(10, 20))

        # create main entry and button
        self.entry = customtkinter.CTkEntry(self, placeholder_text="Input Query", height=50)
        self.entry.grid(row=0, column=1, padx=(20, 0), pady=(20, 0), sticky="ew")

        self.main_button_1 = customtkinter.CTkButton(master=self, fg_color="transparent", border_width=2,
                                                     text_color=("gray10", "#DCE4EE"), width=250, height=50)
        self.main_button_1.grid(row=0, column=2, padx=(20, 20), pady=(20, 0), sticky="ew")

        # =================== Frame 1 Kolom 1 Row 1 ===================
        self.frame1 = customtkinter.CTkFrame(self, fg_color="transparent")
        self.frame1.grid(row=1, column=1, padx=(20, 0), pady=(0, 10), sticky="nsew")
        self.frame1.grid_columnconfigure(0, weight=1)
        self.frame1.grid_rowconfigure(1, weight=1)
        # bikin label
        self.textbox1_label = customtkinter.CTkLabel(self.frame1,
                                                     text="• Menampilkan File yang mirip dengan Query:")
        self.textbox1_label.grid(row=0, column=0, padx=(10, 10), pady=(0, 0), sticky="w")
        # create textbox(ngikut frame 1)
        self.textbox1 = customtkinter.CTkTextbox(self.frame1, width=250)
        self.textbox1.grid(row=1, column=0, padx=0, pady=0, sticky="nsew")

        # ====================== Frame 2 kolom 1 Row 2 ===========================
        self.frame2 = customtkinter.CTkFrame(self, fg_color="transparent")
        self.frame2.grid(row=2, column=1, padx=(20, 0), pady=(0, 20), sticky="nsew")
        self.frame2.grid_columnconfigure(0, weight=1)
        self.frame2.grid_rowconfigure(1, weight=1)

        self.textbox2_label = customtkinter.CTkLabel(self.frame2,
                                                     text="• Menampilkan Pre-Processing, serta kata dasar dan jumlahnya:")
        self.textbox2_label.grid(row=0, column=0, padx=(10, 10), pady=(0, 5), sticky="w")

        self.textbox2 = customtkinter.CTkTextbox(self.frame2, width=250)
        self.textbox2.grid(row=1, column=0, padx=(0, 0), pady=(0, 0), sticky="nsew")

        # create tabview
        self.tabview = customtkinter.CTkTabview(self, width=250)
        self.tabview.grid(row=1, column=2, padx=(20, 20), pady=(2, 20), sticky="nsew")
        # self.tabview.add("CTkTabview")
        self.tabview.add("Anggota")
        self.tabview.add("Saran")
        # self.tabview.tab("CTkTabview").grid_columnconfigure(0, weight=1)  # configure grid of individual tabs
        self.tabview.tab("Anggota").grid_columnconfigure(0, weight=1)
        self.tabview.tab("Saran").grid_columnconfigure(0, weight=1)

        # self.optionmenu_1 = customtkinter.CTkOptionMenu(self.tabview.tab("CTkTabview"), dynamic_resizing=False,
        #                                                 values=["Value 1", "Value 2", "Value Long Long Long"])
        # self.optionmenu_1.grid(row=0, column=0, padx=20, pady=(20, 10))
        # self.combobox_1 = customtkinter.CTkComboBox(self.tabview.tab("CTkTabview"),
        #                                             values=["Value 1", "Value 2", "Value Long....."])
        # self.combobox_1.grid(row=1, column=0, padx=20, pady=(10, 10))
        self.string_input_button = customtkinter.CTkButton(self.tabview.tab("Saran"), text="Pop-Up Kirim Saran",
                                                           command=self.open_input_dialog_event)
        self.string_input_button.grid(row=0, column=0, padx=20, pady=(20, 20))
        self.label_tab_2 = customtkinter.CTkLabel(self.tabview.tab("Anggota"),
                                                  text="Kelompok 8 Kelas BB\n" + "_________________\n\n" + "• [152021082] Rafi P.D\n\n" + "• [152021134] Anggit N.R\n\n" + "• [152021166] M. Ghaza A.L\n\n")
        self.label_tab_2.grid(row=0, column=0, padx=20, pady=20)

        # create scrollable frame
        self.scrollable_frame = customtkinter.CTkScrollableFrame(self,
                                                                 label_text="List Keseluruhan File pada Direktori")
        self.scrollable_frame.grid(row=2, column=2, padx=(20, 20), pady=(10, 20), sticky="nsew")
        self.scrollable_frame.grid_columnconfigure(0, weight=1)
        self.scrollable_frame_switches = []

        # set default values
        self.sidebarMateri.configure(text="Materi", command=self.funcSidebarMateri)
        self.main_button_1.configure(text="Cari", command=self.search_button_event)
        self.appearance_mode_optionemenu.set("Dark")

    def show_popup_textbox(self):
        self.popup = customtkinter.CTkToplevel(self)
        self.popup.title("Kotak Materi")
        # self.popup.geometry("600x400")  # Sesuaikan ukuran sesuai kebutuhan
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        # Mendapatkan ukuran pop-up
        popup_width = 600
        popup_height = 400

        # Menghitung posisi untuk pop-up berada di tengah
        x_position = (screen_width - popup_width) // 2
        y_position = (screen_height - popup_height) // 2

        self.popup.geometry(f"{popup_width}x{popup_height}+{x_position}+{y_position}")

        # Atur penampilan mode "Dark" dan default color theme "blue"
        self.popup.tk_setPalette(background="#263238", foreground="#ffffff", activeBackground="#37474F",
                                 activeForeground="#ffffff")

        # Buat widget Canvas untuk scrollable
        canvas = customtkinter.CTkCanvas(self.popup)
        canvas.pack(side="left", fill="both", expand=True)

        # Buat widget Frame di dalam Canvas
        frame = customtkinter.CTkFrame(canvas)
        canvas.create_window((0, 0), window=frame, anchor="nw")

        # Buat widget teks untuk menampilkan materi di dalam textbox
        materi_text = "Materi tentang Boolean Retrieval\n\n" + "Boolean Retrieval Model (BRM) adalah model Information Retrieval yang menggunakan operator logika AND, OR, dan NOT untuk menghubungkan kata kunci dalam pencarian. Hasilnya berupa nilai binary (1 atau 0), menunjukkan relevansi dokumen dengan query. BRM memeriksa keberadaan atau ketiadaan kata kunci dalam dokumen terindeks, memungkinkan pengguna merinci pencarian dengan operator AND, OR, dan NOT. Dengan demikian, BRM menghasilkan dokumen sesuai kriteria boolean, memudahkan pengguna mendapatkan informasi yang sesuai dengan kebutuhan mereka." + "\n\nRumus dasar untuk Boolean retrieval dapat dijelaskan sebagai berikut : " + "\n1.	AND (dan): Dokumen yang mengandung semua kata kunci yang dicari. Rumusnya adalah AA AND BB, di mana AA dan BB adalah kata kunci" + "\n2.	OR (atau): Dokumen yang mengandung setidaknya satu dari kata kunci yang dicari. Rumusnya adalah AA OR BB." + "\n3.	NOT (bukan): Dokumen yang mengandung kata kunci pertama tetapi tidak mengandung kata kunci kedua. Rumusnya adalah AA NOT BB."
        materi_textbox = customtkinter.CTkTextbox(frame, width=540, height=313)
        materi_textbox.insert("1.0", materi_text)
        materi_textbox.pack(pady=20, padx=20, fill="both",
                            expand=True)  # Atur fill dan expand untuk mengisi seluruh frame

        # Buat tombol untuk menutup pop-up
        close_button = customtkinter.CTkButton(frame, text="Tutup", command=self.popup.destroy)
        close_button.pack(side="bottom", pady=10, padx=10)

        # Tambahkan scrollbar ke Canvas
        scrollbar = tkinter.Scrollbar(self.popup, orient="vertical", command=canvas.yview)
        scrollbar.pack(side="right", fill="y")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Fungsi untuk menyesuaikan ukuran frame saat ukuran canvas berubah
        def configure_frame(event):
            canvas.configure(scrollregion=canvas.bbox("all"))

        canvas.bind("<Configure>", configure_frame)

    def funcSidebarMateri(self):
        print("Button Side Materi")
        self.show_popup_textbox()

    def open_input_dialog_event(self):
        dialog = customtkinter.CTkInputDialog(text="Kirim Saran : ", title="Kotak Saran")
        print("Saran :", dialog.get_input())

    def sidebar_button_event(self):
        print("sidebar_button click")

    def change_appearance_mode_event(self, new_appearance_mode: str):
        customtkinter.set_appearance_mode(new_appearance_mode)

    def search_button_event(self):
        query = self.entry.get()  # Ambil input dari entry
        result_text = self.boolean_search(tokens, query)

        # Tampilkan hasil pada TextBox
        self.textbox1.delete("1.0", tkinter.END)
        self.textbox1.insert(tkinter.END, result_text)

    def clear_scrollable_frame_switches(self):
        for label in self.scrollable_frame_switches:
            label.destroy()
        self.scrollable_frame_switches = []

    def boolean_search(self, tokens, query):
        query_terms = self.stemming(self.filtering(self.tokenizing(self.case_folding(query))))

        print("Pre-processing Query {}:".format(query_terms))
        print('---' * 30)
        result = set()

        if 'and' in query_terms:
            result = set.intersection(*[tokens.get(term, set()) for term in query_terms if term != 'and'])
        elif 'or' in query_terms:
            result = set.union(*[tokens.get(term, set()) for term in query_terms if term != 'or'])
        elif 'not' in query_terms:
            result = set.difference(set(tokens.get(query_terms[0], set())), set(tokens.get(query_terms[2], set())))
        elif 'xor' in query_terms and len(query_terms) == 3:
            # Menangani operasi 'xor' untuk operator 'xor' dengan hanya dua istilah
            term1 = tokens.get(query_terms[0], set())
            term2 = tokens.get(query_terms[2], set())
            result = term1.symmetric_difference(term2)
        else:
            result = set.intersection(*[tokens.get(term, set()) for term in query_terms])


        if not result:
            result_text = "Pre-processing Query {}: \n\nTidak ada dokumen yang sesuai dengan query".format(query_terms)
        else:
            result_text = "Pre-processing Query {}: \n\nFile yang mirip adalah :\n".format(query_terms)

            # Format daftar file yang mirip
            result_files = "- " + "\n- ".join(sorted(result))
            result_text += result_files + "\n"

        return result_text

    def case_folding(self, text):
        return text.lower()

    def tokenizing(self, text):
        return text.split()

    def filtering(self, tokens, custom_stop_words=None):
        stopword_remover = set(custom_stop_words) if custom_stop_words else set()

        def baca_stopwords_docx(file_path):
            doc = docx.Document(file_path)
            stopwords = [paragraph.text for paragraph in doc.paragraphs]
            return stopwords

        daftar_stopwords_docx = baca_stopwords_docx("stopword.docx")
        stopword_remover.update(daftar_stopwords_docx)

        return [token for token in tokens if token not in stopword_remover and token not in string.punctuation]

    def stemming(self, tokens):
        factory = StemmerFactory()
        stemmer = factory.create_stemmer()
        return [stemmer.stem(token) for token in tokens]

    def analisis_file(self, file_path, content):
        content = self.case_folding(content)
        token_tokenizing = self.tokenizing(content)
        token_filter = self.filtering(token_tokenizing)
        tokens = self.stemming(token_filter)

        frekuensi_kata = {}

        nama_file = os.path.basename(file_path)
        ekstensi_file = file_path.split(".")[-1]

        # Menampilkan nama file dan ekstensi pada CTkScrollableFrame
        file_label = customtkinter.CTkLabel(self.scrollable_frame, text=f"{nama_file} ({ekstensi_file})")
        file_label.grid(row=len(self.scrollable_frame_switches), column=0, padx=10, pady=(0, 20))

        self.scrollable_frame_switches.append(file_label)

        # Menampilkan hasil analisis pada textbox2
        self.textbox2.insert(tkinter.END, f"Nama File: {nama_file}\n")
        self.textbox2.insert(tkinter.END, f"Ekstensi file: {ekstensi_file}\n")
        self.textbox2.insert(tkinter.END, "Jenis dan Jumlah kata Dasar dalam file {}:\n".format(nama_file))

        # Menghitung frekuensi kata dasar
        for kata in tokens:
            frekuensi_kata[kata] = frekuensi_kata.get(kata, 0) + 1

        # Menampilkan kata dasar dan jumlahnya di textbox2
        for kata, hitung in frekuensi_kata.items():
            self.textbox2.insert(tkinter.END, "          •[{}]  ===>  {}\n".format(kata, hitung))

        self.textbox2.insert(tkinter.END, "-" * 30 + "\n")

        # self.textbox2.insert(tkinter.END, f"\nContent:\n{content}\n")
        # self.textbox2.insert(tkinter.END, f"\nToken Tokenizing: {token_tokenizing}\n")
        # self.textbox2.insert(tkinter.END, f"\nToken Filter: {token_filter}\n")
        self.textbox2.insert(tkinter.END, f"Hasil Stemming: {tokens}\n")
        self.textbox2.insert(tkinter.END,
                             "______________________________________________________________________________________________________________________" + "\n\n\n")

        return tokens

    def baca_docx(self, file_path):
        doc = docx.Document(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return content

    def baca_pdf(self, file_path):
        with open(file_path, "rb") as file:
            pdf_reader = PdfReader(file)
            content = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                content += page.extract_text()
        return content

    def baca_txt(self, file_path):
        with open(file_path, "r", encoding="utf-8") as file:
            content = file.read()
        return content

    def proses_file_dalam_direktori(self, directory_path):
        self.clear_scrollable_frame_switches()  # Bersihkan list sebelum menambah file baru
        file_contents = {}
        all_tokens = set()

        for nama_file in os.listdir(directory_path):
            file_path = os.path.join(directory_path, nama_file)

            if nama_file.lower().endswith('.docx'):
                content = self.baca_docx(file_path)
                file_contents[os.path.basename(file_path)] = content
                tokens = self.analisis_file(file_path, content)
            elif nama_file.lower().endswith('.pdf'):
                content = self.baca_pdf(file_path)
                file_contents[os.path.basename(file_path)] = content
                tokens = self.analisis_file(file_path, content)
            elif nama_file.lower().endswith('.txt'):
                content = self.baca_txt(file_path)
                file_contents[os.path.basename(file_path)] = content
                tokens = self.analisis_file(file_path, content)

            all_tokens.update(tokens)

        # Convert the set of all tokens to a dictionary with dummy values
        tokens_dict = {term: set() for term in all_tokens}
        for nama_file in os.listdir(directory_path):
            file_path = os.path.join(directory_path, nama_file)

            if nama_file.lower().endswith('.docx'):
                content = self.baca_docx(file_path)
            elif nama_file.lower().endswith('.pdf'):
                content = self.baca_pdf(file_path)
            elif nama_file.lower().endswith('.txt'):
                content = self.baca_txt(file_path)

            tokens = self.analisis_file(file_path, content)

            for term in tokens:
                tokens_dict[term].add(os.path.basename(file_path))

        return file_contents, tokens_dict


if __name__ == "__main__":
    # change the directory path with yours
    directory_path = r"C:\ListOfFile"
    app = App()
    file_contents, tokens = app.proses_file_dalam_direktori(directory_path)
    app.mainloop()
